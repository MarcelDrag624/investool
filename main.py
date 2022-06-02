import requests
from openpyxl import Workbook, load_workbook
import datetime
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

loc = "wallet_data.xlsx"
wb = load_workbook(loc, data_only = True)

document = Document()

class Asset:

    sum_of_invested_money = 0

    profits_and_losses_balance85_prcnt = 0
    only_profits_85prcnt = 0
    only_losses_85prcnt = 0
    money_after_selling_all_85prcnt = 0


    profits_and_losses_balance_98prcnt = 0
    only_profits_98prcnt = 0
    only_losses_98prcnt = 0
    money_after_selling_all_98prcnt = 0

    overall_quantity_of_owned_items = 0

    all_instances = []

    def get_percent_of_money_invested_in_asset(self):
        return self.percent_of_sum_of_invested_money

    def __init__(self, sheet_name):

        self.sheet = wb[sheet_name]
        self.asset_name = self.sheet['B1'].value
        self.asset_url = self.sheet['A1'].value
        self.asset_owned_quantity = self.sheet['B15'].value
        self.asset_invested_money = self.sheet['C15'].value
        self.asset_average_money_spent_on_single_piece = self.asset_invested_money / int(self.asset_owned_quantity)
        self.asset_data_from_api = requests.get(self.asset_url).json()
        self.asset_current_lowest_price = float(self.asset_data_from_api['lowest_price'].replace(',',".").replace('zł',""))
        self.asset_current_median_price = float(self.asset_data_from_api['median_price'].replace(',',".").replace('zł',""))
        self.asset_quantity_sold_in_last_24h = self.asset_data_from_api['volume']
        self.balance85prcnt = (self.asset_owned_quantity * self.asset_current_median_price * 0.85) - (self.asset_invested_money)
        self.balance98prcnt = (self.asset_owned_quantity * self.asset_current_median_price * 0.98) - (self.asset_invested_money)
        Asset.sum_of_invested_money += self.asset_invested_money
        Asset.profits_and_losses_balance85_prcnt += self.balance85prcnt

        if self.balance85prcnt < 0:
            Asset.only_losses_85prcnt += self.balance85prcnt
        elif self.balance85prcnt > 0:
            Asset.only_profits_85prcnt += self.balance85prcnt

        Asset.profits_and_losses_balance_98prcnt += self.balance98prcnt
        if self.balance98prcnt < 0:
            Asset.only_losses_98prcnt += self.balance98prcnt
        elif self.balance98prcnt > 0:
            Asset.only_profits_98prcnt += self.balance98prcnt

        self.balance85prcnt_in_percents = (self.balance85prcnt/self.asset_invested_money)*100
        self.balance98prcnt_in_percents = (self.balance98prcnt/self.asset_invested_money)*100

        Asset.overall_quantity_of_owned_items += self.asset_owned_quantity

        Asset.all_instances.append(self)


    def calc_percent_of_money_invested_in_asset(self):
        self.percent_of_sum_of_invested_money = (self.asset_invested_money/Asset.sum_of_invested_money)*100


    def print_asset_summary(self):

        print(self.asset_name)
        print(f'Lowest price on Steam Market: {self.asset_current_lowest_price} zł.')
        print(f'Median price on Steam Market: {self.asset_current_median_price} zł.')
        print(f'Items sold in the last 24h: {self.asset_quantity_sold_in_last_24h}.')
        print(f'Average price paid for a single piece: {round(self.asset_average_money_spent_on_single_piece,2)} zł.')
        print(f'Owned quantity: {round(self.asset_owned_quantity,2)}.')
        print(f'Money invested in this asset: {round(self.asset_invested_money,2)} zł.')
        print(f'Balance: {round(self.balance85prcnt,2)} zł (after 15% Steam Market commission) or {round(self.balance98prcnt,2)} zł (after 2% csdeals.com commission).')
        print(f'Balance in percents: {round(self.balance85prcnt_in_percents,2)}% (after 15% Steam Market commission) or {round(self.balance98prcnt_in_percents,2)}% (after 2% csdeals.com commission).')
        print(f'Percent of all money invested in this wallet: {round(self.percent_of_sum_of_invested_money,2)}%.\n')


    def add_asset_summary_to_report(self):

        global document

        p0 = document.add_paragraph()
        p0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        title = p0.add_run(f'{self.asset_name}:')
        # title_font = title.font
        title.font.size = Pt(16)
        title.font.bold = True

        p1 = document.add_paragraph(style = 'List Bullet')
        style = document.styles['Normal']
        style.paragraph_format.line_spacing = 1
        p1.add_run(f'Lowest').bold = True
        p1.add_run(f' price on Steam Market: ')
        p1.add_run(f'{self.asset_current_lowest_price} zł').bold = True
        p1.add_run(f'.')

        p2 = document.add_paragraph(style = 'List Bullet')
        p2.add_run(f'Median').bold = True
        p2.add_run(f' price on Steam Market: ')
        p2.add_run(f'{self.asset_current_median_price} zł').bold = True
        p2.add_run(f'.')

        p3 = document.add_paragraph(style = 'List Bullet')
        p3.add_run(f'Items sold ').bold = True
        p3.add_run(f'in the ')
        p3.add_run(f'last 24h').bold = True
        p3.add_run(f': ')
        p3.add_run(f'{self.asset_quantity_sold_in_last_24h}').bold = True
        p3.add_run(f'.')

        p4 = document.add_paragraph(style = 'List Bullet')
        p4.add_run(f'Average price paid').bold = True
        p4.add_run(f' for a single piece: ')
        p4.add_run(f'{round(self.asset_average_money_spent_on_single_piece,2)} zł').bold = True
        p4.add_run(f'.')

        p5 = document.add_paragraph(style = 'List Bullet')
        p5.add_run(f'Quantity ').bold = True
        p5.add_run(f'of items ')
        p5.add_run(f'owned: ').bold = True
        p5.add_run(f'{self.asset_owned_quantity}').bold = True
        p5.add_run(f'.')

        p6 = document.add_paragraph(style = 'List Bullet')
        p6.add_run(f'Money invested ').bold = True
        p6.add_run(f'in this asset: ')
        p6.add_run(f'{round(self.asset_invested_money,2)} zł').bold = True
        p6.add_run(f'.')

        p7 = document.add_paragraph(style = 'List Bullet')
        p7.add_run(f'Balance: {round(self.balance85prcnt,2)} zł ').bold = True
        p7.add_run(f'(after 15% Steam Market commission) or ')
        p7.add_run(f'{round(self.balance98prcnt,2)} zł ').bold = True
        p7.add_run(f'(after 2% csdeals.com commission).')

        p8 = document.add_paragraph(style = 'List Bullet')
        p8.add_run(f'Balance in percents: {round(self.balance85prcnt_in_percents,2)}% ').bold = True
        p8.add_run(f'(after 15% Steam Market commission) or ')
        p8.add_run(f'{round(self.balance98prcnt_in_percents,2)}% ').bold = True
        p8.add_run(f'(after 2% csdeals.com commission).')

        p9 = document.add_paragraph(style = 'List Bullet')
        p9.add_run(f'Percent of all money invested ').bold = True
        p9.add_run(f'in this wallet: ')
        p9.add_run(f'{round(self.percent_of_sum_of_invested_money,2)}%').bold = True
        p9.add_run(f'.\n')


    def print_general_summary():
        Asset.money_after_selling_all_85prcnt = Asset.profits_and_losses_balance85_prcnt + Asset.sum_of_invested_money
        Asset.money_after_selling_all_98prcnt = Asset.profits_and_losses_balance_98prcnt + Asset.sum_of_invested_money
        print(f'Money invested in this wallet: {round(Asset.sum_of_invested_money,2)} zł.')
        print(f'Overall quantity of owned items: {Asset.overall_quantity_of_owned_items}.')
        print(f'Wallet\'s total value: {round(Asset.money_after_selling_all_85prcnt,2)} zł (after 15% commission - Steam Market) or {round(Asset.money_after_selling_all_98prcnt,2)} zł (after 2% commission - csdeals.com).')
        print(f'Wallet\'s balance: {round(Asset.profits_and_losses_balance85_prcnt,2)} zł (after 15% commission - Steam Market) or {round(Asset.profits_and_losses_balance_98prcnt,2)} zł (after 2% commission - csdeals.com).')
        print(f'Balance after selling only profitable assets: {round(Asset.only_profits_85prcnt,2)} zł (after 15% commission - Steam Market) or {round(Asset.only_profits_98prcnt,2)} zł (after 2% commission - csdeals.com).')
        print(f'Balance after selling only lossable assets: {round(Asset.only_losses_85prcnt,2)} zł (after 15% commission - Steam Market) or {round(Asset.only_losses_98prcnt,2)} zł (after 2% commission - csdeals.com).\n')


    def add_general_summary_to_report():
        Asset.money_after_selling_all_85prcnt = Asset.profits_and_losses_balance85_prcnt + Asset.sum_of_invested_money
        Asset.money_after_selling_all_98prcnt = Asset.profits_and_losses_balance_98prcnt + Asset.sum_of_invested_money

        p0 = document.add_paragraph()
        p0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # style = document.styles['Normal']
        # style.paragraph_format.line_spacing = 1.5
        title = p0.add_run(f'General summary:')
        # title_font = title.font
        title.font.size = Pt(16)
        title.font.bold = True

        p1 = document.add_paragraph(style = 'List Bullet')
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p1.add_run(f'Money invested ').bold = True
        p1.add_run(f'in this ')
        p1.add_run(f'wallet: ').bold = True
        p1.add_run(f'{round(Asset.sum_of_invested_money,2)} zł').bold = True
        p1.add_run(f'.')

        p2 = document.add_paragraph(style = 'List Bullet')
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p2.add_run(f'Overall quantity ').bold = True
        p2.add_run(f'of owned ')
        p2.add_run(f'items: ').bold = True
        p2.add_run(f'{Asset.overall_quantity_of_owned_items}').bold = True
        p2.add_run(f'.')

        p3 = document.add_paragraph(style = 'List Bullet')
        p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p3.add_run(f'Wallet\'s total value: ').bold = True
        p3.add_run(f'{round(Asset.money_after_selling_all_85prcnt,2)} zł ').bold = True
        p3.add_run(f'(after 15% Steam Market commission) or ')
        p3.add_run(f'{round(Asset.money_after_selling_all_98prcnt,2)} zł ').bold = True
        p3.add_run(f'(after 2% csdeals.com commission).')

        p4 = document.add_paragraph(style = 'List Bullet')
        p4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p4.add_run(f'Wallet\'s balance: ').bold = True
        p4.add_run(f'{round(Asset.profits_and_losses_balance85_prcnt,2)} zł ').bold = True
        p4.add_run(f'(after 15% Steam Market commission) or ')
        p4.add_run(f'{round(Asset.profits_and_losses_balance_98prcnt,2)} zł ').bold = True
        p4.add_run(f'(after 2% csdeals.com commission).')

        p5 = document.add_paragraph(style = 'List Bullet')
        p5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p5.add_run(f'Balance ').bold = True
        p5.add_run(f'after selling ')
        p5.add_run(f'only profitable ').bold = True
        p5.add_run(f'assets: ')
        p5.add_run(f'{round(Asset.only_profits_85prcnt,2)} zł ').bold = True
        p5.add_run(f'(after 15% Steam Market commission) or ')
        p5.add_run(f'{round(Asset.only_profits_98prcnt,2)} zł ').bold = True
        p5.add_run(f'(after 2% csdeals.com commission).')

        p6 = document.add_paragraph(style = 'List Bullet')
        p6.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p6.add_run(f'Balance ').bold = True
        p6.add_run(f'after selling ')
        p6.add_run(f'only lossable ').bold = True
        p6.add_run(f'assets: ')
        p6.add_run(f'{round(Asset.only_losses_85prcnt,2)} zł ').bold = True
        p6.add_run(f'(after 15% Steam Market commission) or ')
        p6.add_run(f'{round(Asset.only_losses_98prcnt,2)} zł ').bold = True
        p6.add_run(f'(after 2% csdeals.com commission).')


for i in range(0,len(wb.sheetnames)-2):
    Asset(wb.sheetnames[i])

for i in range(0,len(Asset.all_instances)):
    Asset.all_instances[i].calc_percent_of_money_invested_in_asset()

Asset.all_instances.sort(key=Asset.get_percent_of_money_invested_in_asset)
Asset.all_instances.reverse()

for i in range(0,len(Asset.all_instances)):
    Asset.all_instances[i].add_asset_summary_to_report()

Asset.add_general_summary_to_report()

document.save('report.docx')
